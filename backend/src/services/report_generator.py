import asyncio
import json
import os
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import tempfile
import io

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import LineChart, BarChart, Reference
from openpyxl.utils.dataframe import dataframe_to_rows

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus import PageBreak, Image
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib import rcParams
import structlog

from src.core.config import settings
from src.models.inspection import InspectionStatus, CheckItemStatus

logger = structlog.get_logger()

# 设置中文字体支持
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

class ReportGenerator:
    """巡检报告生成器"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "inspection_reports"
        self.temp_dir.mkdir(exist_ok=True)
        
        # 报告样式配置
        self.styles = {
            'title_font': Font(name='Arial', size=16, bold=True),
            'header_font': Font(name='Arial', size=12, bold=True),
            'normal_font': Font(name='Arial', size=10),
            'success_fill': PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid'),
            'warning_fill': PatternFill(start_color='FFD700', end_color='FFD700', fill_type='solid'),
            'error_fill': PatternFill(start_color='FFB6C1', end_color='FFB6C1', fill_type='solid'),
            'header_fill': PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        }
        
    async def generate_inspection_report(
        self,
        inspection_data: Dict[str, Any],
        format_type: str = "excel",
        include_charts: bool = True
    ) -> str:
        """生成巡检报告

        Args:
            inspection_data: 巡检数据
            format_type: 报告格式 ("excel", "pdf", "html", "word")
            include_charts: 是否包含图表

        Returns:
            生成的报告文件路径
        """
        try:
            format_lower = format_type.lower()
            if format_lower == "excel":
                return await self._generate_excel_report(inspection_data, include_charts)
            elif format_lower == "pdf":
                return await self._generate_pdf_report(inspection_data, include_charts)
            elif format_lower == "html":
                return await self._generate_html_report(inspection_data, include_charts)
            elif format_lower == "word":
                return await self._generate_word_report(inspection_data, include_charts)
            else:
                raise ValueError(f"不支持的报告格式: {format_type}. 支持的格式: excel, pdf, html, word")

        except Exception as e:
            logger.error("生成巡检报告失败", error=str(e), format=format_type)
            raise
    
    async def _generate_excel_report(
        self,
        inspection_data: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """生成Excel格式报告"""
        try:
            # 创建工作簿
            wb = Workbook()
            
            # 删除默认工作表
            wb.remove(wb.active)
            
            # 创建各个工作表
            await self._create_summary_sheet(wb, inspection_data)
            await self._create_device_details_sheet(wb, inspection_data)
            await self._create_check_results_sheet(wb, inspection_data)
            
            if include_charts:
                await self._create_charts_sheet(wb, inspection_data)
            
            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"inspection_report_{timestamp}.xlsx"
            file_path = self.temp_dir / filename
            
            # 保存工作簿
            wb.save(str(file_path))
            
            logger.info("Excel报告生成成功", file_path=str(file_path))
            return str(file_path)
            
        except Exception as e:
            logger.error("生成Excel报告失败", error=str(e))
            raise
    
    async def _create_summary_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建汇总工作表"""
        ws = wb.create_sheet("巡检汇总")
        
        # 设置标题
        ws.merge_cells('A1:F1')
        ws['A1'] = f"网络设备巡检报告 - {data.get('inspection_name', '未命名巡检')}"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # 基本信息
        row = 3
        basic_info = [
            ("巡检ID:", data.get('inspection_id', 'N/A')),
            ("巡检时间:", data.get('inspection_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))),
            ("设备总数:", len(data.get('devices', []))),
            ("巡检状态:", data.get('status', InspectionStatus.UNKNOWN)),
            ("执行时长:", f"{data.get('execution_duration', 0)} 秒"),
        ]
        
        for label, value in basic_info:
            ws[f'A{row}'] = label
            ws[f'A{row}'].font = self.styles['header_font']
            ws[f'B{row}'] = str(value)
            row += 1
        
        # 统计摘要
        row += 2
        ws[f'A{row}'] = "检查统计摘要"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1
        
        summary_stats = data.get('summary_stats', {})
        stats_info = [
            ("总检查项数:", summary_stats.get('total_checks', 0)),
            ("通过检查项:", summary_stats.get('passed_checks', 0)),
            ("失败检查项:", summary_stats.get('failed_checks', 0)),
            ("警告检查项:", summary_stats.get('warning_checks', 0)),
            ("错误检查项:", summary_stats.get('error_checks', 0)),
            ("通过率:", f"{summary_stats.get('pass_rate', 0):.1f}%"),
        ]
        
        for label, value in stats_info:
            ws[f'A{row}'] = label
            ws[f'A{row}'].font = self.styles['header_font']
            ws[f'B{row}'] = str(value)
            row += 1
        
        # 设备状态汇总表
        row += 2
        ws[f'A{row}'] = "设备状态汇总"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1
        
        # 表头
        headers = ['设备名称', '设备IP', '设备类型', '巡检状态', '通过率', '问题数量']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.styles['header_font']
            cell.fill = self.styles['header_fill']
            cell.font = Font(color='FFFFFF', bold=True)
        
        row += 1
        
        # 设备数据
        devices = data.get('devices', [])
        for device in devices:
            device_data = [
                device.get('device_name', 'N/A'),
                device.get('ip_address', 'N/A'),
                device.get('device_type', 'N/A'),
                device.get('inspection_status', 'Unknown'),
                f"{device.get('pass_rate', 0):.1f}%",
                device.get('issue_count', 0)
            ]
            
            for col, value in enumerate(device_data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                # 根据状态设置颜色
                if col == 4:  # 巡检状态列
                    if 'success' in str(value).lower() or 'completed' in str(value).lower():
                        cell.fill = self.styles['success_fill']
                    elif 'warning' in str(value).lower():
                        cell.fill = self.styles['warning_fill']
                    elif 'fail' in str(value).lower() or 'error' in str(value).lower():
                        cell.fill = self.styles['error_fill']
            row += 1
        
        # 调整列宽
        column_widths = [15, 15, 12, 12, 10, 10]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width
    
    async def _create_device_details_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建设备详情工作表"""
        ws = wb.create_sheet("设备详情")
        
        # 标题
        ws.merge_cells('A1:H1')
        ws['A1'] = "设备详细信息"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')
        
        row = 3
        devices = data.get('devices', [])
        
        for device in devices:
            # 设备基本信息
            ws[f'A{row}'] = f"设备: {device.get('device_name', 'Unknown')}"
            ws[f'A{row}'].font = self.styles['header_font']
            row += 1
            
            device_info = [
                ("IP地址:", device.get('ip_address', 'N/A')),
                ("设备类型:", device.get('device_type', 'N/A')),
                ("厂商:", device.get('vendor', 'N/A')),
                ("型号:", device.get('model', 'N/A')),
                ("系统版本:", device.get('software_version', 'N/A')),
                ("运行时间:", device.get('uptime', 'N/A')),
                ("最后巡检时间:", device.get('last_inspection', 'N/A')),
            ]
            
            for label, value in device_info:
                ws[f'A{row}'] = label
                ws[f'B{row}'] = str(value)
                row += 1
            
            # 性能指标
            performance = device.get('performance_metrics', {})
            if performance:
                row += 1
                ws[f'A{row}'] = "性能指标"
                ws[f'A{row}'].font = self.styles['header_font']
                row += 1
                
                perf_metrics = [
                    ("CPU使用率:", f"{performance.get('cpu_usage', 'N/A')}%"),
                    ("内存使用率:", f"{performance.get('memory_usage', 'N/A')}%"),
                    ("活跃接口数:", performance.get('active_interfaces', 'N/A')),
                    ("接口总数:", performance.get('total_interfaces', 'N/A')),
                ]
                
                for label, value in perf_metrics:
                    ws[f'A{row}'] = label
                    ws[f'B{row}'] = str(value)
                    row += 1
            
            row += 2  # 设备间距
        
        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 25
    
    async def _create_check_results_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建检查结果工作表"""
        ws = wb.create_sheet("检查结果详情")
        
        # 标题
        ws.merge_cells('A1:G1')
        ws['A1'] = "检查结果详情"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # 表头
        row = 3
        headers = ['设备名称', '检查项', '检查类型', '状态', '期望值', '实际值', '执行时间(ms)']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = self.styles['header_font']
            cell.fill = self.styles['header_fill']
            cell.font = Font(color='FFFFFF', bold=True)
        
        row += 1
        
        # 检查结果数据
        devices = data.get('devices', [])
        for device in devices:
            device_name = device.get('device_name', 'Unknown')
            check_results = device.get('check_results', [])
            
            for result in check_results:
                result_data = [
                    device_name,
                    result.get('check_item_name', 'N/A'),
                    result.get('check_item_type', 'N/A'),
                    result.get('status', 'Unknown'),
                    result.get('expected_value', 'N/A'),
                    result.get('actual_value', 'N/A'),
                    result.get('execution_time', 0)
                ]
                
                for col, value in enumerate(result_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    
                    # 根据状态设置颜色
                    if col == 4:  # 状态列
                        status = str(value).lower()
                        if status == CheckItemStatus.PASS.lower():
                            cell.fill = self.styles['success_fill']
                        elif status == CheckItemStatus.WARNING.lower():
                            cell.fill = self.styles['warning_fill']
                        elif status in [CheckItemStatus.FAIL.lower(), CheckItemStatus.ERROR.lower()]:
                            cell.fill = self.styles['error_fill']
                
                row += 1
        
        # 调整列宽
        column_widths = [15, 20, 12, 10, 15, 15, 12]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width
    
    async def _create_charts_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建图表工作表"""
        ws = wb.create_sheet("图表分析")
        
        # 标题
        ws.merge_cells('A1:F1')
        ws['A1'] = "巡检结果图表分析"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # 创建状态分布图表数据
        row = 3
        ws[f'A{row}'] = "检查结果状态分布"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1
        
        # 状态统计数据
        summary_stats = data.get('summary_stats', {})
        status_data = [
            ["状态", "数量"],
            ["通过", summary_stats.get('passed_checks', 0)],
            ["失败", summary_stats.get('failed_checks', 0)],
            ["警告", summary_stats.get('warning_checks', 0)],
            ["错误", summary_stats.get('error_checks', 0)]
        ]
        
        for row_data in status_data:
            for col, value in enumerate(row_data, 1):
                ws.cell(row=row, column=col, value=value)
            row += 1
        
        # 创建柱状图
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "检查结果状态分布"
        chart.y_axis.title = '数量'
        chart.x_axis.title = '状态'
        
        data_ref = Reference(ws, min_col=2, min_row=5, max_row=8, max_col=2)
        cats = Reference(ws, min_col=1, min_row=6, max_row=8)
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats)
        
        ws.add_chart(chart, "D4")
    
    async def _generate_pdf_report(
        self,
        inspection_data: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """生成PDF格式报告"""
        try:
            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"inspection_report_{timestamp}.pdf"
            file_path = self.temp_dir / filename
            
            # 创建PDF文档
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12
            )
            
            # 构建PDF内容
            story = []
            
            # 标题
            title = Paragraph(
                f"网络设备巡检报告<br/>{inspection_data.get('inspection_name', '未命名巡检')}",
                title_style
            )
            story.append(title)
            story.append(Spacer(1, 20))
            
            # 基本信息
            story.append(Paragraph("巡检基本信息", heading_style))
            
            basic_info_data = [
                ['项目', '值'],
                ['巡检ID', inspection_data.get('inspection_id', 'N/A')],
                ['巡检时间', inspection_data.get('inspection_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))],
                ['设备总数', str(len(inspection_data.get('devices', [])))],
                ['巡检状态', inspection_data.get('status', InspectionStatus.UNKNOWN)],
                ['执行时长', f"{inspection_data.get('execution_duration', 0)} 秒"],
            ]
            
            basic_info_table = Table(basic_info_data, colWidths=[2*inch, 3*inch])
            basic_info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(basic_info_table)
            story.append(Spacer(1, 20))
            
            # 统计摘要
            story.append(Paragraph("检查统计摘要", heading_style))
            
            summary_stats = inspection_data.get('summary_stats', {})
            summary_data = [
                ['统计项', '数量'],
                ['总检查项数', str(summary_stats.get('total_checks', 0))],
                ['通过检查项', str(summary_stats.get('passed_checks', 0))],
                ['失败检查项', str(summary_stats.get('failed_checks', 0))],
                ['警告检查项', str(summary_stats.get('warning_checks', 0))],
                ['错误检查项', str(summary_stats.get('error_checks', 0))],
                ['通过率', f"{summary_stats.get('pass_rate', 0):.1f}%"],
            ]
            
            summary_table = Table(summary_data, colWidths=[2*inch, 2*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(PageBreak())
            
            # 设备详情
            story.append(Paragraph("设备巡检详情", heading_style))
            
            devices = inspection_data.get('devices', [])
            for device in devices:
                # 设备标题
                device_title = Paragraph(
                    f"设备: {device.get('device_name', 'Unknown')}",
                    heading_style
                )
                story.append(device_title)
                
                # 设备信息表
                device_data = [
                    ['属性', '值'],
                    ['IP地址', device.get('ip_address', 'N/A')],
                    ['设备类型', device.get('device_type', 'N/A')],
                    ['厂商', device.get('vendor', 'N/A')],
                    ['巡检状态', device.get('inspection_status', 'Unknown')],
                    ['通过率', f"{device.get('pass_rate', 0):.1f}%"],
                ]
                
                device_table = Table(device_data, colWidths=[2*inch, 3*inch])
                device_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(device_table)
                story.append(Spacer(1, 15))
            
            # 生成PDF
            doc.build(story)
            
            logger.info("PDF报告生成成功", file_path=str(file_path))
            return str(file_path)
            
        except Exception as e:
            logger.error("生成PDF报告失败", error=str(e))
            raise
    
    async def generate_batch_report(
        self,
        inspection_list: List[Dict[str, Any]],
        format_type: str = "excel"
    ) -> str:
        """生成批量巡检报告"""
        try:
            # 合并多次巡检数据
            merged_data = {
                'inspection_name': f"批量巡检报告 ({len(inspection_list)} 次巡检)",
                'inspection_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'devices': [],
                'summary_stats': {
                    'total_checks': 0,
                    'passed_checks': 0,
                    'failed_checks': 0,
                    'warning_checks': 0,
                    'error_checks': 0,
                    'pass_rate': 0
                }
            }
            
            # 合并统计数据
            all_devices = []
            for inspection in inspection_list:
                devices = inspection.get('devices', [])
                all_devices.extend(devices)
                
                stats = inspection.get('summary_stats', {})
                merged_data['summary_stats']['total_checks'] += stats.get('total_checks', 0)
                merged_data['summary_stats']['passed_checks'] += stats.get('passed_checks', 0)
                merged_data['summary_stats']['failed_checks'] += stats.get('failed_checks', 0)
                merged_data['summary_stats']['warning_checks'] += stats.get('warning_checks', 0)
                merged_data['summary_stats']['error_checks'] += stats.get('error_checks', 0)
            
            # 计算总通过率
            total_checks = merged_data['summary_stats']['total_checks']
            if total_checks > 0:
                passed_checks = merged_data['summary_stats']['passed_checks']
                merged_data['summary_stats']['pass_rate'] = (passed_checks / total_checks) * 100
            
            merged_data['devices'] = all_devices
            
            # 生成报告
            return await self.generate_inspection_report(merged_data, format_type)

        except Exception as e:
            logger.error("生成批量报告失败", error=str(e))
            raise

    async def _generate_html_report(
        self,
        inspection_data: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """生成HTML格式报告"""
        try:
            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"inspection_report_{timestamp}.html"
            file_path = self.temp_dir / filename

            # 构建HTML内容
            html_content = self._build_html_content(inspection_data, include_charts)

            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info("HTML报告生成成功", file_path=str(file_path))
            return str(file_path)

        except Exception as e:
            logger.error("生成HTML报告失败", error=str(e))
            raise

    def _build_html_content(self, data: Dict[str, Any], include_charts: bool) -> str:
        """构建HTML内容"""
        summary_stats = data.get('summary_stats', {})
        devices = data.get('devices', [])

        # 计算通过率颜色
        pass_rate = summary_stats.get('pass_rate', 0)
        pass_rate_color = '#4caf50' if pass_rate >= 90 else '#ff9800' if pass_rate >= 70 else '#f44336'

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>巡检报告 - {data.get('inspection_name', '未命名巡检')}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            background: #f5f7fa;
            padding: 20px;
            line-height: 1.6;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            text-align: center;
            padding-bottom: 10px;
            border-bottom: 3px solid #3498db;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #34495e;
            margin: 30px 0 15px 0;
            padding-bottom: 8px;
            border-bottom: 2px solid #ecf0f1;
        }}
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 15px;
            margin-bottom: 25px;
        }}
        .info-item {{
            background: #f8f9fa;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid #3498db;
        }}
        .info-label {{
            font-weight: bold;
            color: #7f8c8d;
            font-size: 0.9em;
            margin-bottom: 5px;
        }}
        .info-value {{
            color: #2c3e50;
            font-size: 1.1em;
        }}
        .stats-cards {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 15px;
            margin: 20px 0;
        }}
        .stat-card {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .stat-card.success {{ background: linear-gradient(135deg, #4caf50 0%, #45a049 100%); }}
        .stat-card.warning {{ background: linear-gradient(135deg, #ff9800 0%, #f57c00 100%); }}
        .stat-card.error {{ background: linear-gradient(135deg, #f44336 0%, #e53935 100%); }}
        .stat-card .value {{
            font-size: 2.5em;
            font-weight: bold;
            margin: 10px 0;
        }}
        .stat-card .label {{
            font-size: 0.9em;
            opacity: 0.9;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ecf0f1;
        }}
        th {{
            background: #3498db;
            color: white;
            font-weight: 600;
        }}
        tr:hover {{
            background: #f8f9fa;
        }}
        .status-badge {{
            padding: 4px 12px;
            border-radius: 12px;
            font-size: 0.85em;
            font-weight: 600;
        }}
        .status-online {{ background: #d4edda; color: #155724; }}
        .status-warning {{ background: #fff3cd; color: #856404; }}
        .status-error {{ background: #f8d7da; color: #721c24; }}
        .status-offline {{ background: #e2e3e5; color: #383d41; }}
        .progress-bar {{
            width: 100%;
            height: 24px;
            background: #ecf0f1;
            border-radius: 12px;
            overflow: hidden;
            margin: 10px 0;
        }}
        .progress-fill {{
            height: 100%;
            background: {pass_rate_color};
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-weight: bold;
            font-size: 0.9em;
            transition: width 0.3s ease;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ecf0f1;
            color: #7f8c8d;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>🔍 网络设备巡检报告</h1>
        <p style="text-align: center; color: #7f8c8d; margin-bottom: 30px;">
            {data.get('inspection_name', '未命名巡检')}
        </p>
        <h2>📋 基本信息</h2>
        <div class="info-grid">
            <div class="info-item">
                <div class="info-label">巡检ID</div>
                <div class="info-value">{data.get('inspection_id', 'N/A')}</div>
            </div>
            <div class="info-item">
                <div class="info-label">巡检时间</div>
                <div class="info-value">{data.get('inspection_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}</div>
            </div>
            <div class="info-item">
                <div class="info-label">设备总数</div>
                <div class="info-value">{len(devices)}</div>
            </div>
            <div class="info-item">
                <div class="info-label">执行时长</div>
                <div class="info-value">{data.get('execution_duration', 0)} 秒</div>
            </div>
        </div>
        <h2>📊 统计摘要</h2>
        <div class="stats-cards">
            <div class="stat-card">
                <div class="label">总检查项</div>
                <div class="value">{summary_stats.get('total_checks', 0)}</div>
            </div>
            <div class="stat-card success">
                <div class="label">通过</div>
                <div class="value">{summary_stats.get('passed_checks', 0)}</div>
            </div>
            <div class="stat-card warning">
                <div class="label">警告</div>
                <div class="value">{summary_stats.get('warning_checks', 0)}</div>
            </div>
            <div class="stat-card error">
                <div class="label">失败</div>
                <div class="value">{summary_stats.get('failed_checks', 0)}</div>
            </div>
        </div>
        <div class="progress-bar">
            <div class="progress-fill" style="width: {pass_rate:.1f}%;">
                通过率: {pass_rate:.1f}%
            </div>
        </div>
        <h2>🖥️ 设备详情</h2>
        <table>
            <thead>
                <tr>
                    <th>设备名称</th>
                    <th>IP地址</th>
                    <th>设备类型</th>
                    <th>状态</th>
                    <th>通过率</th>
                    <th>问题数量</th>
                </tr>
            </thead>
            <tbody>
"""

        # 添加设备行
        for device in devices:
            device_name = device.get('device_name', 'Unknown')
            ip_address = device.get('ip_address', 'N/A')
            device_type = device.get('device_type', 'N/A')
            inspection_status = device.get('inspection_status', 'Unknown')
            device_pass_rate = device.get('pass_rate', 0)
            issue_count = device.get('issue_count', 0)

            # 确定状态样式
            status_class = 'status-offline'
            if 'success' in inspection_status.lower() or 'completed' in inspection_status.lower():
                status_class = 'status-online'
            elif 'warning' in inspection_status.lower():
                status_class = 'status-warning'
            elif 'fail' in inspection_status.lower() or 'error' in inspection_status.lower():
                status_class = 'status-error'

            html += f"""
                <tr>
                    <td><strong>{device_name}</strong></td>
                    <td>{ip_address}</td>
                    <td>{device_type}</td>
                    <td><span class="status-badge {status_class}">{inspection_status}</span></td>
                    <td>{device_pass_rate:.1f}%</td>
                    <td>{issue_count}</td>
                </tr>
"""

        html += """
            </tbody>
        </table>
        <div class="footer">
            <p>此报告由网络设备巡检系统自动生成</p>
            <p>生成时间: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + """</p>
        </div>
    </div>
</body>
</html>
"""

        return html

    async def _generate_word_report(
        self,
        inspection_data: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """生成Word格式报告"""
        try:
            # 注意：需要安装python-docx库
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                error_msg = "生成Word报告需要python-docx库，请运行: uv pip install python-docx"
                logger.error(error_msg)
                raise ImportError(error_msg)

            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"inspection_report_{timestamp}.docx"
            file_path = self.temp_dir / filename

            # 创建文档
            doc = Document()

            # 设置页面属性
            section = doc.sections[0]
            section.page_height = Inches(11.69)  # A4高度
            section.page_width = Inches(8.27)    # A4宽度

            # 标题
            title = doc.add_heading('网络设备巡检报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 副标题
            subtitle = doc.add_paragraph(inspection_data.get('inspection_name', '未命名巡检'))
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(14)
            subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()  # 空行

            # 基本信息
            doc.add_heading('基本信息', 1)

            basic_info_table = doc.add_table(rows=5, cols=2)
            basic_info_table.style = 'Light Grid Accent 1'

            basic_info_data = [
                ('巡检ID', inspection_data.get('inspection_id', 'N/A')),
                ('巡检时间', inspection_data.get('inspection_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))),
                ('设备总数', str(len(inspection_data.get('devices', [])))),
                ('巡检状态', inspection_data.get('status', 'Unknown')),
                ('执行时长', f"{inspection_data.get('execution_duration', 0)} 秒")
            ]

            for idx, (label, value) in enumerate(basic_info_data):
                row = basic_info_table.rows[idx]
                row.cells[0].text = label
                row.cells[1].text = str(value)
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()  # 空行

            # 统计摘要
            doc.add_heading('统计摘要', 1)

            summary_stats = inspection_data.get('summary_stats', {})
            summary_table = doc.add_table(rows=6, cols=2)
            summary_table.style = 'Light Grid Accent 1'

            summary_data = [
                ('总检查项数', str(summary_stats.get('total_checks', 0))),
                ('通过检查项', str(summary_stats.get('passed_checks', 0))),
                ('失败检查项', str(summary_stats.get('failed_checks', 0))),
                ('警告检查项', str(summary_stats.get('warning_checks', 0))),
                ('错误检查项', str(summary_stats.get('error_checks', 0))),
                ('通过率', f"{summary_stats.get('pass_rate', 0):.1f}%")
            ]

            for idx, (label, value) in enumerate(summary_data):
                row = summary_table.rows[idx]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_page_break()

            # 设备详情
            doc.add_heading('设备巡检详情', 1)

            devices = inspection_data.get('devices', [])
            if devices:
                device_table = doc.add_table(rows=1, cols=6)
                device_table.style = 'Light Grid Accent 1'

                # 表头
                header_cells = device_table.rows[0].cells
                headers = ['设备名称', 'IP地址', '设备类型', '巡检状态', '通过率', '问题数']
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    header_cells[idx].paragraphs[0].runs[0].font.bold = True

                # 数据行
                for device in devices:
                    row_cells = device_table.add_row().cells
                    row_cells[0].text = device.get('device_name', 'Unknown')
                    row_cells[1].text = device.get('ip_address', 'N/A')
                    row_cells[2].text = device.get('device_type', 'N/A')
                    row_cells[3].text = device.get('inspection_status', 'Unknown')
                    row_cells[4].text = f"{device.get('pass_rate', 0):.1f}%"
                    row_cells[5].text = str(device.get('issue_count', 0))

            # 页脚
            doc.add_paragraph()
            footer = doc.add_paragraph('此报告由网络设备巡检系统自动生成')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(10)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            time_footer = doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            time_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_footer.runs[0].font.size = Pt(10)
            time_footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # 保存文档
            doc.save(str(file_path))

            logger.info("Word报告生成成功", file_path=str(file_path))
            return str(file_path)

        except ImportError as e:
            error_msg = "生成Word报告需要python-docx库，请运行: uv pip install python-docx"
            logger.error(error_msg, error=str(e))
            raise ImportError(error_msg)
        except Exception as e:
            logger.error("生成Word报告失败", error=str(e))
            raise

    async def cleanup_temp_files(self, older_than_hours: int = 24):
        """清理临时文件"""
        try:
            cutoff_time = datetime.now() - timedelta(hours=older_than_hours)
            
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_time < cutoff_time:
                        file_path.unlink()
                        logger.info("清理临时报告文件", file_path=str(file_path))
                        
        except Exception as e:
            logger.warning("清理临时文件失败", error=str(e))

# 全局报告生成器实例
report_generator = ReportGenerator()