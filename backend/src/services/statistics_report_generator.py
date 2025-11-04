"""
统计报表生成器
专门用于生成统计报表（Excel、PDF、HTML、Word格式）
"""
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
import tempfile
import io

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.drawing.image import Image as XLImage

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus import PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端
import structlog

logger = structlog.get_logger()

# 设置中文字体支持
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


class StatisticsReportGenerator:
    """统计报表生成器"""

    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "statistics_reports"
        self.temp_dir.mkdir(exist_ok=True)

        # 报告样式配置
        self.styles = {
            'title_font': Font(name='Arial', size=16, bold=True),
            'header_font': Font(name='Arial', size=12, bold=True),
            'normal_font': Font(name='Arial', size=10),
            'header_fill': PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        }

    async def generate_statistics_report(
        self,
        statistics_data: Dict[str, Any],
        title: str = "统计报表",
        format_type: str = "excel",
        include_charts: bool = True
    ) -> str:
        """生成统计报表

        Args:
            statistics_data: 统计数据（从statistics_service获取）
            title: 报表标题
            format_type: 报告格式 ("excel", "pdf", "html", "word")
            include_charts: 是否包含图表

        Returns:
            生成的报告文件路径
        """
        try:
            format_lower = format_type.lower()
            if format_lower == "excel":
                return await self._generate_excel_report(statistics_data, title, include_charts)
            elif format_lower == "pdf":
                return await self._generate_pdf_report(statistics_data, title, include_charts)
            elif format_lower == "html":
                return await self._generate_html_report(statistics_data, title, include_charts)
            elif format_lower == "word":
                return await self._generate_word_report(statistics_data, title, include_charts)
            else:
                raise ValueError(f"不支持的报告格式: {format_type}")

        except Exception as e:
            logger.error("生成统计报表失败", error=str(e), format=format_type)
            raise

    async def _generate_excel_report(
        self,
        data: Dict[str, Any],
        title: str,
        include_charts: bool = True
    ) -> str:
        """生成Excel格式统计报表"""
        try:
            wb = Workbook()
            wb.remove(wb.active)

            # 创建各个工作表
            await self._create_overview_sheet(wb, data, title)
            await self._create_device_distribution_sheet(wb, data)
            await self._create_performance_sheet(wb, data)
            await self._create_top_devices_sheet(wb, data)

            if include_charts:
                await self._create_charts_sheet(wb, data)

            # 生成文件路径
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"statistics_report_{timestamp}.xlsx"
            file_path = self.temp_dir / filename

            wb.save(str(file_path))

            logger.info("Excel统计报表生成成功", file_path=str(file_path))
            return str(file_path)

        except Exception as e:
            logger.error("生成Excel统计报表失败", error=str(e))
            raise

    async def _create_overview_sheet(self, wb: Workbook, data: Dict[str, Any], title: str):
        """创建概览工作表"""
        ws = wb.create_sheet("统计概览")

        # 标题
        ws.merge_cells('A1:D1')
        ws['A1'] = title
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')

        row = 3
        overview = data.get('overview', {})

        # 基本统计
        ws[f'A{row}'] = "基本统计"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1

        basic_stats = [
            ("设备总数", overview.get('total_devices', 0)),
            ("在线设备", overview.get('active_devices', 0)),
            ("离线设备", overview.get('offline_devices', 0)),
            ("告警设备", overview.get('warning_devices', 0)),
            ("故障设备", overview.get('error_devices', 0)),
            ("平均正常运行时间", f"{overview.get('avg_uptime', 0):.1f}小时"),
            ("总巡检次数", overview.get('total_executions', 0)),
            ("平均健康评分", f"{overview.get('avg_score', 0):.1f}"),
        ]

        for label, value in basic_stats:
            ws[f'A{row}'] = label
            ws[f'A{row}'].font = self.styles['header_font']
            ws[f'B{row}'] = str(value)
            row += 1

        # 调整列宽
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 20

    async def _create_device_distribution_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建设备分布工作表"""
        ws = wb.create_sheet("设备分布")

        # 标题
        ws.merge_cells('A1:C1')
        ws['A1'] = "设备分布统计"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')

        row = 3
        distribution = data.get('device_distribution', {})

        # 按类型分布
        ws[f'A{row}'] = "按类型分布"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1

        headers = ['设备类型', '数量', '占比']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(color='FFFFFF', bold=True)
            cell.fill = self.styles['header_fill']
        row += 1

        by_type = distribution.get('by_type', {})
        total = sum(by_type.values()) if by_type else 1

        for device_type, count in by_type.items():
            ws[f'A{row}'] = device_type
            ws[f'B{row}'] = count
            ws[f'C{row}'] = f"{(count / total * 100):.1f}%"
            row += 1

        row += 2

        # 按位置分布
        ws[f'A{row}'] = "按位置分布"
        ws[f'A{row}'].font = self.styles['header_font']
        row += 1

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(color='FFFFFF', bold=True)
            cell.fill = self.styles['header_fill']
        row += 1

        by_location = distribution.get('by_location', {})
        total_loc = sum(by_location.values()) if by_location else 1

        for location, count in by_location.items():
            ws[f'A{row}'] = location
            ws[f'B{row}'] = count
            ws[f'C{row}'] = f"{(count / total_loc * 100):.1f}%"
            row += 1

        # 调整列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15

    async def _create_performance_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建性能统计工作表"""
        ws = wb.create_sheet("性能统计")

        # 标题
        ws.merge_cells('A1:E1')
        ws['A1'] = "设备性能统计"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')

        row = 3
        headers = ['设备名称', 'CPU使用率', '内存使用率', '可用性', '健康评分']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(color='FFFFFF', bold=True)
            cell.fill = self.styles['header_fill']
        row += 1

        perf_stats = data.get('performance_stats', {})
        by_device = perf_stats.get('by_device', [])

        for device in by_device:
            metrics = device.get('metrics', {})
            ws[f'A{row}'] = device.get('device_name', 'N/A')
            ws[f'B{row}'] = f"{metrics.get('cpu_usage', 0):.1f}%"
            ws[f'C{row}'] = f"{metrics.get('memory_usage', 0):.1f}%"
            ws[f'D{row}'] = f"{metrics.get('availability', 0):.1f}%"
            ws[f'E{row}'] = f"{metrics.get('health_score', 0):.1f}"
            row += 1

        # 调整列宽
        for i in range(1, 6):
            ws.column_dimensions[chr(64 + i)].width = 18

    async def _create_top_devices_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建TOP设备工作表"""
        ws = wb.create_sheet("TOP设备")

        # 标题
        ws.merge_cells('A1:D1')
        ws['A1'] = "TOP 10 性能设备"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')

        row = 3
        headers = ['排名', '设备名称', '设备类型', '性能评分']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(color='FFFFFF', bold=True)
            cell.fill = self.styles['header_fill']
        row += 1

        top_devices = data.get('top_devices', {}).get('by_performance', [])

        for idx, device in enumerate(top_devices[:10], 1):
            ws[f'A{row}'] = idx
            ws[f'B{row}'] = device.get('device_name', 'N/A')
            ws[f'C{row}'] = device.get('device_type', 'N/A')
            ws[f'D{row}'] = f"{device.get('score', 0):.2f}"
            row += 1

        # 调整列宽
        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15

    async def _create_charts_sheet(self, wb: Workbook, data: Dict[str, Any]):
        """创建图表工作表"""
        ws = wb.create_sheet("数据图表")

        # 标题
        ws.merge_cells('A1:F1')
        ws['A1'] = "统计图表"
        ws['A1'].font = self.styles['title_font']
        ws['A1'].alignment = Alignment(horizontal='center')

        row = 3

        # 设备类型分布数据
        distribution = data.get('device_distribution', {})
        by_type = distribution.get('by_type', {})

        if by_type:
            ws[f'A{row}'] = "设备类型分布"
            ws[f'A{row}'].font = self.styles['header_font']
            row += 1

            ws[f'A{row}'] = "类型"
            ws[f'B{row}'] = "数量"
            row += 1

            chart_start_row = row
            for device_type, count in by_type.items():
                ws[f'A{row}'] = device_type
                ws[f'B{row}'] = count
                row += 1

            # 创建柱状图
            chart = BarChart()
            chart.title = "设备类型分布"
            chart.style = 10
            chart.y_axis.title = '数量'
            chart.x_axis.title = '设备类型'

            data_ref = Reference(ws, min_col=2, min_row=chart_start_row, max_row=row-1)
            cats = Reference(ws, min_col=1, min_row=chart_start_row+1, max_row=row-1)
            chart.add_data(data_ref, titles_from_data=False)
            chart.set_categories(cats)

            ws.add_chart(chart, f"D{chart_start_row}")

    async def _generate_pdf_report(
        self,
        data: Dict[str, Any],
        title: str,
        include_charts: bool = True
    ) -> str:
        """生成PDF格式统计报表"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"statistics_report_{timestamp}.pdf"
            file_path = self.temp_dir / filename

            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            styles = getSampleStyleSheet()

            # 自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )

            story = []

            # 标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))

            # 概览数据
            overview = data.get('overview', {})
            overview_data = [
                ['统计项', '数值'],
                ['设备总数', str(overview.get('total_devices', 0))],
                ['在线设备', str(overview.get('active_devices', 0))],
                ['离线设备', str(overview.get('offline_devices', 0))],
                ['平均健康评分', f"{overview.get('avg_score', 0):.1f}"],
            ]

            overview_table = Table(overview_data, colWidths=[3*inch, 2*inch])
            overview_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(overview_table)
            story.append(PageBreak())

            # 设备分布
            story.append(Paragraph("设备分布统计", styles['Heading2']))

            distribution = data.get('device_distribution', {})
            by_type = distribution.get('by_type', {})

            if by_type:
                type_data = [['设备类型', '数量']]
                for device_type, count in by_type.items():
                    type_data.append([device_type, str(count)])

                type_table = Table(type_data, colWidths=[3*inch, 2*inch])
                type_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(type_table)

            # 生成PDF
            doc.build(story)

            logger.info("PDF统计报表生成成功", file_path=str(file_path))
            return str(file_path)

        except Exception as e:
            logger.error("生成PDF统计报表失败", error=str(e))
            raise

    async def _generate_html_report(
        self,
        data: Dict[str, Any],
        title: str,
        include_charts: bool = True
    ) -> str:
        """生成HTML格式统计报表"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"statistics_report_{timestamp}.html"
            file_path = self.temp_dir / filename

            html_content = self._build_statistics_html(data, title, include_charts)

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info("HTML统计报表生成成功", file_path=str(file_path))
            return str(file_path)

        except Exception as e:
            logger.error("生成HTML统计报表失败", error=str(e))
            raise

    def _build_statistics_html(self, data: Dict[str, Any], title: str, include_charts: bool) -> str:
        """构建HTML统计报表内容"""
        overview = data.get('overview', {})
        distribution = data.get('device_distribution', {})

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            background: #f5f7fa;
            padding: 20px;
            line-height: 1.6;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            text-align: center;
            padding-bottom: 10px;
            border-bottom: 3px solid #3498db;
            margin-bottom: 30px;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 30px 0;
        }}
        .stat-card {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
        }}
        .stat-card .value {{
            font-size: 2em;
            font-weight: bold;
            margin: 10px 0;
        }}
        .stat-card .label {{
            font-size: 0.9em;
            opacity: 0.9;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ecf0f1;
        }}
        th {{
            background: #3498db;
            color: white;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            color: #7f8c8d;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{title}</h1>

        <div class="stats-grid">
            <div class="stat-card">
                <div class="label">设备总数</div>
                <div class="value">{overview.get('total_devices', 0)}</div>
            </div>
            <div class="stat-card">
                <div class="label">在线设备</div>
                <div class="value">{overview.get('active_devices', 0)}</div>
            </div>
            <div class="stat-card">
                <div class="label">离线设备</div>
                <div class="value">{overview.get('offline_devices', 0)}</div>
            </div>
            <div class="stat-card">
                <div class="label">平均评分</div>
                <div class="value">{overview.get('avg_score', 0):.1f}</div>
            </div>
        </div>

        <h2>设备类型分布</h2>
        <table>
            <thead>
                <tr>
                    <th>设备类型</th>
                    <th>数量</th>
                    <th>占比</th>
                </tr>
            </thead>
            <tbody>
"""

        by_type = distribution.get('by_type', {})
        total = sum(by_type.values()) if by_type else 1

        for device_type, count in by_type.items():
            percentage = (count / total * 100) if total > 0 else 0
            html += f"""
                <tr>
                    <td>{device_type}</td>
                    <td>{count}</td>
                    <td>{percentage:.1f}%</td>
                </tr>
"""

        html += """
            </tbody>
        </table>

        <div class="footer">
            <p>此报告由网络设备巡检系统自动生成</p>
            <p>生成时间: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + """</p>
        </div>
    </div>
</body>
</html>
"""

        return html

    async def _generate_word_report(
        self,
        data: Dict[str, Any],
        title: str,
        include_charts: bool = True
    ) -> str:
        """生成Word格式统计报表"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            error_msg = "生成Word报告需要python-docx库"
            logger.error(error_msg)
            raise ImportError(error_msg)

        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"statistics_report_{timestamp}.docx"
            file_path = self.temp_dir / filename

            doc = Document()

            # 标题
            doc_title = doc.add_heading(title, 0)
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # 概览统计
            doc.add_heading('统计概览', 1)

            overview = data.get('overview', {})
            overview_table = doc.add_table(rows=5, cols=2)
            overview_table.style = 'Light Grid Accent 1'

            overview_data = [
                ('设备总数', str(overview.get('total_devices', 0))),
                ('在线设备', str(overview.get('active_devices', 0))),
                ('离线设备', str(overview.get('offline_devices', 0))),
                ('平均健康评分', f"{overview.get('avg_score', 0):.1f}"),
                ('总巡检次数', str(overview.get('total_executions', 0)))
            ]

            for idx, (label, value) in enumerate(overview_data):
                row = overview_table.rows[idx]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_page_break()

            # 设备分布
            doc.add_heading('设备分布', 1)

            distribution = data.get('device_distribution', {})
            by_type = distribution.get('by_type', {})

            if by_type:
                dist_table = doc.add_table(rows=1, cols=2)
                dist_table.style = 'Light Grid Accent 1'

                # 表头
                header_cells = dist_table.rows[0].cells
                header_cells[0].text = '设备类型'
                header_cells[1].text = '数量'

                # 数据
                for device_type, count in by_type.items():
                    row_cells = dist_table.add_row().cells
                    row_cells[0].text = device_type
                    row_cells[1].text = str(count)

            # 页脚
            doc.add_paragraph()
            footer = doc.add_paragraph('此报告由网络设备巡检系统自动生成')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(10)

            # 保存文档
            doc.save(str(file_path))

            logger.info("Word统计报表生成成功", file_path=str(file_path))
            return str(file_path)

        except ImportError as e:
            raise
        except Exception as e:
            logger.error("生成Word统计报表失败", error=str(e))
            raise


# 全局实例
statistics_report_generator = StatisticsReportGenerator()
