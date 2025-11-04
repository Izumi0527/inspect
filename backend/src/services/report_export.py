"""
报表导出功能
支持PDF和Word格式的报表生成和导出
"""
import asyncio
import os
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from pathlib import Path
import tempfile
import uuid

# PDF生成库
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor

# Word文档生成库
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

import structlog

logger = structlog.get_logger()

class ReportExporter:
    \"\"\"报表导出服务\"\"\"
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / \"inspect_reports\"
        self.temp_dir.mkdir(exist_ok=True)
        
        # 报表模板配置
        self.report_templates = {
            'device_summary': '设备汇总报表',
            'inspection_report': '巡检结果报表', 
            'alert_report': '告警统计报表',
            'performance_report': '性能分析报表'
        }
    
    async def generate_pdf_report(
        self, 
        report_type: str,
        data: Dict[str, Any],
        title: str = None,
        subtitle: str = None
    ) -> str:
        \"\"\"生成PDF报表\"\"\"
        try:
            report_id = str(uuid.uuid4())
            filename = f\"{report_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{report_id[:8]}.pdf\"
            filepath = self.temp_dir / filename
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # 构建报表内容
            story = []
            styles = getSampleStyleSheet()
            
            # 添加标题
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                textColor=HexColor('#2563eb'),
                alignment=1  # 居中
            )
            
            report_title = title or self.report_templates.get(report_type, '系统报表')
            story.append(Paragraph(report_title, title_style))
            
            # 添加副标题
            if subtitle:
                subtitle_style = ParagraphStyle(
                    'CustomSubtitle',
                    parent=styles['Heading2'],
                    fontSize=12,
                    spaceAfter=20,
                    textColor=HexColor('#64748b'),
                    alignment=1
                )
                story.append(Paragraph(subtitle, subtitle_style))
            
            # 添加生成时间
            time_style = ParagraphStyle(
                'TimeStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor('#6b7280'),
                alignment=2  # 右对齐
            )
            story.append(Paragraph(f\"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\", time_style))
            story.append(Spacer(1, 20))
            
            # 根据报表类型生成内容
            if report_type == 'device_summary':
                story.extend(self._build_device_summary_pdf(data, styles))
            elif report_type == 'inspection_report':
                story.extend(self._build_inspection_report_pdf(data, styles))
            elif report_type == 'alert_report':
                story.extend(self._build_alert_report_pdf(data, styles))
            elif report_type == 'performance_report':
                story.extend(self._build_performance_report_pdf(data, styles))
            else:
                # 通用报表格式
                story.extend(self._build_generic_pdf(data, styles))
            
            # 生成PDF
            doc.build(story)
            
            logger.info(\"PDF report generated\", 
                       report_type=report_type,
                       filename=filename,
                       filepath=str(filepath))
            
            return str(filepath)
            
        except Exception as e:
            logger.error(\"Failed to generate PDF report\",
                        report_type=report_type,
                        error=str(e))
            raise e
    
    async def generate_word_report(
        self,
        report_type: str,
        data: Dict[str, Any],
        title: str = None,
        subtitle: str = None
    ) -> str:
        \"\"\"生成Word报表\"\"\"
        try:
            report_id = str(uuid.uuid4())
            filename = f\"{report_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{report_id[:8]}.docx\"
            filepath = self.temp_dir / filename
            
            # 创建Word文档
            doc = Document()
            
            # 设置文档标题
            report_title = title or self.report_templates.get(report_type, '系统报表')
            title_paragraph = doc.add_heading(report_title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加副标题
            if subtitle:
                subtitle_paragraph = doc.add_heading(subtitle, level=2)
                subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加生成信息
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run(f\"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\")
            info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加分隔线
            doc.add_paragraph(\"_\" * 50)
            
            # 根据报表类型生成内容
            if report_type == 'device_summary':
                self._build_device_summary_word(doc, data)
            elif report_type == 'inspection_report':
                self._build_inspection_report_word(doc, data)
            elif report_type == 'alert_report':
                self._build_alert_report_word(doc, data)
            elif report_type == 'performance_report':
                self._build_performance_report_word(doc, data)
            else:
                # 通用报表格式
                self._build_generic_word(doc, data)
            
            # 保存文档
            doc.save(str(filepath))
            
            logger.info(\"Word report generated\",
                       report_type=report_type, 
                       filename=filename,
                       filepath=str(filepath))
            
            return str(filepath)
            
        except Exception as e:
            logger.error(\"Failed to generate Word report\",
                        report_type=report_type,
                        error=str(e))
            raise e
    
    def _build_device_summary_pdf(self, data: Dict[str, Any], styles) -> List:
        \"\"\"构建设备汇总PDF内容\"\"\"
        story = []
        
        # 概览统计
        story.append(Paragraph(\"设备概览统计\", styles['Heading2']))
        
        summary_data = [
            ['统计项', '数值', '占比'],
            ['设备总数', str(data.get('total', 0)), '100%'],
            ['在线设备', str(data.get('online', 0)), f\"{data.get('online', 0) / max(data.get('total', 1), 1) * 100:.1f}%\"],
            ['离线设备', str(data.get('offline', 0)), f\"{data.get('offline', 0) / max(data.get('total', 1), 1) * 100:.1f}%\"],
            ['告警设备', str(data.get('warning', 0)), f\"{data.get('warning', 0) / max(data.get('total', 1), 1) * 100:.1f}%\"]
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 1*inch, 1*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # 设备列表
        if 'devices' in data and data['devices']:
            story.append(Paragraph(\"设备详情列表\", styles['Heading2']))
            
            device_data = [['设备名称', 'IP地址', '设备类型', '状态', '位置']]
            for device in data['devices'][:20]:  # 限制显示20个设备
                device_data.append([
                    device.get('name', ''),
                    device.get('ip', ''),
                    device.get('device_type', ''),
                    device.get('status', ''),
                    device.get('location', '')
                ])
            
            device_table = Table(device_data, colWidths=[1.5*inch, 1.2*inch, 1*inch, 0.8*inch, 1.5*inch])
            device_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#2563eb')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(device_table)
        
        return story
    
    def _build_device_summary_word(self, doc: Document, data: Dict[str, Any]):
        \"\"\"构建设备汇总Word内容\"\"\"
        # 概览统计
        doc.add_heading('设备概览统计', level=1)
        
        # 创建统计表格
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '统计项'
        hdr_cells[1].text = '数值'
        hdr_cells[2].text = '占比'
        
        # 数据行
        rows_data = [
            ('设备总数', str(data.get('total', 0)), '100%'),
            ('在线设备', str(data.get('online', 0)), f\"{data.get('online', 0) / max(data.get('total', 1), 1) * 100:.1f}%\"),
            ('离线设备', str(data.get('offline', 0)), f\"{data.get('offline', 0) / max(data.get('total', 1), 1) * 100:.1f}%\"),
            ('告警设备', str(data.get('warning', 0)), f\"{data.get('warning', 0) / max(data.get('total', 1), 1) * 100:.1f}%\")
        ]
        
        for i, row_data in enumerate(rows_data, 1):
            cells = table.rows[i].cells
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]
            cells[2].text = row_data[2]
        
        # 设备列表
        if 'devices' in data and data['devices']:
            doc.add_page_break()
            doc.add_heading('设备详情列表', level=1)
            
            # 创建设备表格
            device_table = doc.add_table(rows=1, cols=5)
            device_table.style = 'Table Grid'
            
            # 表头
            hdr_cells = device_table.rows[0].cells
            hdr_cells[0].text = '设备名称'
            hdr_cells[1].text = 'IP地址'
            hdr_cells[2].text = '设备类型'
            hdr_cells[3].text = '状态'
            hdr_cells[4].text = '位置'
            
            # 设备数据
            for device in data['devices'][:50]:  # 限制显示50个设备
                row_cells = device_table.add_row().cells
                row_cells[0].text = device.get('name', '')
                row_cells[1].text = device.get('ip', '')
                row_cells[2].text = device.get('device_type', '')
                row_cells[3].text = device.get('status', '')
                row_cells[4].text = device.get('location', '')
    
    def _build_generic_pdf(self, data: Dict[str, Any], styles) -> List:
        \"\"\"构建通用PDF内容\"\"\"
        story = []
        
        for key, value in data.items():
            story.append(Paragraph(f\"<b>{key}:</b> {str(value)}\", styles['Normal']))
            story.append(Spacer(1, 12))
        
        return story
    
    def _build_generic_word(self, doc: Document, data: Dict[str, Any]):
        \"\"\"构建通用Word内容\"\"\"
        for key, value in data.items():
            p = doc.add_paragraph()
            p.add_run(f\"{key}: \").bold = True
            p.add_run(str(value))
    
    def _build_inspection_report_pdf(self, data: Dict[str, Any], styles) -> List:
        \"\"\"构建巡检报告PDF内容\"\"\"
        # TODO: 实现巡检报告PDF格式
        return self._build_generic_pdf(data, styles)
    
    def _build_inspection_report_word(self, doc: Document, data: Dict[str, Any]):
        \"\"\"构建巡检报告Word内容\"\"\"
        # TODO: 实现巡检报告Word格式
        self._build_generic_word(doc, data)
    
    def _build_alert_report_pdf(self, data: Dict[str, Any], styles) -> List:
        \"\"\"构建告警报告PDF内容\"\"\"
        # TODO: 实现告警报告PDF格式
        return self._build_generic_pdf(data, styles)
    
    def _build_alert_report_word(self, doc: Document, data: Dict[str, Any]):
        \"\"\"构建告警报告Word内容\"\"\"
        # TODO: 实现告警报告Word格式
        self._build_generic_word(doc, data)
    
    def _build_performance_report_pdf(self, data: Dict[str, Any], styles) -> List:
        \"\"\"构建性能报告PDF内容\"\"\"
        # TODO: 实现性能报告PDF格式
        return self._build_generic_pdf(data, styles)
    
    def _build_performance_report_word(self, doc: Document, data: Dict[str, Any]):
        \"\"\"构建性能报告Word内容\"\"\"
        # TODO: 实现性能报告Word格式
        self._build_generic_word(doc, data)
    
    async def cleanup_old_reports(self, hours: int = 24):
        \"\"\"清理过期的临时报表文件\"\"\"
        try:
            cutoff_time = datetime.now() - timedelta(hours=hours)
            cleaned_count = 0
            
            for file_path in self.temp_dir.glob(\"*\"):
                if file_path.is_file():
                    file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_mtime < cutoff_time:
                        file_path.unlink()
                        cleaned_count += 1
            
            logger.info(\"Old reports cleaned up\",
                       cleaned_count=cleaned_count,
                       hours=hours)
            
            return cleaned_count
            
        except Exception as e:
            logger.error(\"Failed to cleanup old reports\",
                        error=str(e))
            return 0


# 创建全局实例
report_exporter = ReportExporter()